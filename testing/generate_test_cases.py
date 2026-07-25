#!/usr/bin/env python3
"""Generate Test Case Suite for FreeLedger Contract Lifecycle"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_para(text):
    doc.add_paragraph(text)

def add_tc_table(rows):
    """Add a test case as a formatted table"""
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    # Header
    table.rows[0].cells[0].text = 'Field'
    table.rows[0].cells[1].text = 'Value'
    for i, (field, value) in enumerate(rows):
        table.rows[i+1].cells[0].text = field
        table.rows[i+1].cells[1].text = str(value)
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Pt(120)
        row.cells[1].width = Pt(400)
    doc.add_paragraph()  # spacing

# ====== TITLE ======
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Test Case Suite — FreeLedger Contract Lifecycle')
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()
add_para(f'Document Version: 1.0')
add_para(f'Date: {datetime.date.today().strftime("%d %B %Y")}')
add_para('Scope: 35 test cases covering Smart Contract Engine (M2), Workflow API (M5),\n'
         'Wallet Auth (M1), Security, and Consistency across the blockchain contract lifecycle.')
add_para('Derived from: RTP.txt (90-test-case inventory), mapped to 38 requirements.')
doc.add_page_break()

# ====== TEST CASES ======

# ---- M1: Wallet Auth ----
add_heading('MODULE M1 — WALLET AUTHENTICATION', 2)

test_cases_m1 = [
    {
        'id': 'TC-M1-01',
        'title': 'MetaMask Auth Success',
        'req': 'FR-2, FR-6',
        'priority': 'High',
        'technique': 'EP — Valid partition',
        'precondition': 'MetaMask installed and unlocked; backend running on port 8000',
        'steps': '1. User clicks "Connect Wallet"\n2. MetaMask popup appears\n3. User signs EIP-712 message\n4. Frontend sends signature to /auth/verify\n5. Backend verifies and creates session',
        'data': 'Valid Ethereum address (checksum); valid EIP-712 signature',
        'expected': 'HTTP 201; session token returned; JSON with role and expiry',
        'passfail': 'PASS if 201 + session created + role correct',
    },
    {
        'id': 'TC-M1-02',
        'title': 'MetaMask Auth Rejection',
        'req': 'FR-2, FR-6',
        'priority': 'High',
        'technique': 'EP — Invalid partition',
        'precondition': 'MetaMask installed; user at login screen',
        'steps': '1. User clicks "Connect Wallet"\n2. MetaMask popup appears\n3. User clicks "Reject"\n4. Observe frontend',
        'data': 'Valid address; user declines signing',
        'expected': 'No session created; frontend shows error "Signature rejected"',
        'passfail': 'PASS if no session and error shown',
    },
]

for tc in test_cases_m1:
    add_para(f'TC ID: {tc["id"]} — {tc["title"]}')
    add_tc_table([
        ('Req Mapping', tc['req']),
        ('Priority', tc['priority']),
        ('Technique', tc['technique']),
        ('Precondition', tc['precondition']),
        ('Test Steps', tc['steps']),
        ('Test Data', tc['data']),
        ('Expected Result', tc['expected']),
        ('Pass/Fail Criteria', tc['passfail']),
    ])

# ---- M2: Smart Contract Engine ----
add_heading('MODULE M2 — SMART CONTRACT ENGINE', 2)

test_cases_m2 = [
    {
        'id': 'TC-M2-01',
        'title': 'Contract Deployment Success',
        'req': 'FR-1',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Hardhat compiled and running; deployer account funded',
        'steps': '1. Run deploy script (npx hardhat run scripts/deploy.js)\n2. Wait for receipt\n3. Query contract address and state',
        'data': 'Valid Solidity source; deployer private key with test ETH',
        'expected': 'Contract deployed; address non-zero; state == Created (0)',
        'passfail': 'PASS if contract address returned and state == 0',
    },
    {
        'id': 'TC-M2-02',
        'title': 'contractExists() — Valid ID returns true',
        'req': 'FR-1, FR-16',
        'priority': 'High',
        'technique': 'BVA — ID boundary (0, 1, 999)',
        'precondition': 'Contract deployed with ID = 1',
        'steps': '1. Call contractExists(0) → expect false\n2. Call contractExists(1) → expect true\n3. Call contractExists(999) → expect false',
        'data': 'IDs 0, 1, 999',
        'expected': 'contractExists(0) = false, contractExists(1) = true, contractExists(999) = false',
        'passfail': 'PASS if mapping correct; non-existent IDs return false',
    },
    {
        'id': 'TC-M2-03',
        'title': 'setFreelancer() — Client Sets Freelancer',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Contract deployed (ID=1); client wallet known; freelancer wallet known',
        'steps': '1. Client calls setFreelancer(1, freelancer_addr)\n2. Await receipt\n3. Call getContractDetails(1)',
        'data': 'Contract ID = 1, freelancer = 0x70997970...',
        'expected': 'Transaction succeeds; getContractDetails shows freelancer set',
        'passfail': 'PASS if freelancer address matches after call',
    },
    {
        'id': 'TC-M2-04',
        'title': 'setFreelancer() — Non-Client Rejected',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'White Box — require() false branch (onlyClient)',
        'precondition': 'Contract deployed; freelancer wallet',
        'steps': '1. Freelancer calls setFreelancer(1, freelancer_addr)\n2. Observe revert',
        'data': 'Caller = freelancer wallet; Contract ID = 1',
        'expected': 'Transaction reverts with "Only client"',
        'passfail': 'PASS if revert message includes "Only client"',
    },
    {
        'id': 'TC-M2-05',
        'title': 'setFreelancer() — Zero Address Rejected',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'BVA — Invalid boundary',
        'precondition': 'Contract deployed; client wallet',
        'steps': '1. Client calls setFreelancer(1, address(0))\n2. Observe revert',
        'data': 'address(0) = 0x0000000000000000000000000000000000000000',
        'expected': 'Transaction reverts; require(freelancer != address(0)) enforced',
        'passfail': 'PASS if revert prevents zero-address assignment',
    },
    {
        'id': 'TC-M2-06',
        'title': 'setFreelancer() — Non-Existent Contract Rejected',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'Negative Testing',
        'precondition': 'Contract ID 1 exists; ID 999 does not',
        'steps': '1. Client calls setFreelancer(999, freelancer_addr)\n2. Observe revert',
        'data': 'Contract ID = 999 (non-existent)',
        'expected': 'Revert with "Contract does not exist"',
        'passfail': 'PASS if revert message indicates non-existent contract',
    },
    {
        'id': 'TC-M2-07',
        'title': 'fundContract() — Client Funds Escrow',
        'req': 'FR-1, FR-16',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Contract deployed; freelancer set; client wallet funded',
        'steps': '1. Client calls fundContract(1) with value = totalAmount\n2. Await receipt\n3. Check contract status and balance',
        'data': 'Value = 200 ETH; Contract ID = 1',
        'expected': 'Status transitions from 0 (Created) to 1 (InProgress); balance == funded amount',
        'passfail': 'PASS if status == 1 and contract.balance == amount',
    },
    {
        'id': 'TC-M2-08',
        'title': 'fundContract() — Non-Client Caller Rejected',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'White Box — require() false (onlyClient)',
        'precondition': 'Contract deployed; random wallet (not client)',
        'steps': '1. Random wallet calls fundContract(1)\n2. Observe revert',
        'data': 'Caller = random wallet (0x3C44...); value = 1 wei',
        'expected': 'Revert with "Only client"',
        'passfail': 'PASS if revert',
    },
    {
        'id': 'TC-M2-09',
        'title': 'submitMilestone() — Freelancer Submits Work',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Contract funded (InProgress); freelancer wallet',
        'steps': '1. Freelancer calls submitMilestone(1, 0, "QmTestDeliverable")\n2. Await receipt\n3. Check milestone status via getMilestoneDetails',
        'data': 'Contract ID = 1, Milestone Index = 0, CID = "QmTestDeliverable"',
        'expected': 'Milestone status transitions from Pending(1) to Submitted(2)',
        'passfail': 'PASS if milestone.status == 2 (Submitted)',
    },
    {
        'id': 'TC-M2-10',
        'title': 'approveMilestone() — Client Approves, ETH Released',
        'req': 'FR-10, FR-16',
        'priority': 'High',
        'technique': 'Integration',
        'precondition': 'Milestone 0 in Submitted state; client wallet; contract funded',
        'steps': '1. Client calls approveMilestone(1, 0)\n2. Await receipt\n3. Check milestone status and freelancer balance',
        'data': 'Contract ID = 1, Milestone Index = 0',
        'expected': 'Milestone status = 3 (Approved); freelancer receives payout - 2.5% fee',
        'passfail': 'PASS if milestone approved and freelancer balance increased by (amount - fee)',
    },
    {
        'id': 'TC-M2-11',
        'title': 'approveMilestone() — Double Approve Rejected',
        'req': 'FR-10, NFR-8',
        'priority': 'High',
        'technique': 'DT / Security',
        'precondition': 'Milestone 0 already Approved',
        'steps': '1. Client calls approveMilestone(1, 0) again\n2. Observe revert',
        'data': 'Milestone already in Approved state',
        'expected': 'Revert; no second transfer; state remains Approved',
        'passfail': 'PASS if second call blocked (reentrancy/double-payment prevention)',
    },
    {
        'id': 'TC-M2-12',
        'title': 'getContractBalance() — Returns Locked Funds',
        'req': 'FR-1',
        'priority': 'Medium',
        'technique': 'Functional',
        'precondition': 'Contract funded with 200 ETH',
        'steps': '1. Call getContractBalance(1)\n2. Compare with expected',
        'data': 'Contract ID = 1',
        'expected': 'Returns correct amount (200 ETH in wei)',
        'passfail': 'PASS if returned value matches funded amount minus released payments',
    },
    {
        'id': 'TC-M2-13',
        'title': 'Platform Fee — 2.5% Deduction Correct',
        'req': 'FR-16',
        'priority': 'High',
        'technique': 'White Box — Calculation verification',
        'precondition': 'PLATFORM_FEE_BPS = 250 (2.5%)',
        'steps': '1. Call PLATFORM_FEE_BPS()\n2. Compute fee = milestoneAmount * 250 / 10000\n3. Verify payout = milestoneAmount - fee',
        'data': 'Milestone amount = 200 ETH = 200000000000000000000 wei',
        'expected': 'Fee = 5 ETH; Payout = 195 ETH; Platform fee stored correctly',
        'passfail': 'PASS if fee = 2.5% and payout is correct',
    },
    {
        'id': 'TC-M2-14',
        'title': 'Event Logging — Update Events Emitted',
        'req': 'FR-1, NFR-8',
        'priority': 'Medium',
        'technique': 'White Box — Event coverage',
        'precondition': 'Contract deployed; node accessible',
        'steps': '1. Trigger each state transition\n2. Query event logs by topic\n3. Inspect newState and value fields',
        'data': 'Full lifecycle (setFreelancer → fundContract → submitMilestone → approveMilestone)',
        'expected': 'Events emitted with correct newState values and proper ordering',
        'passfail': 'PASS if events present with correct params',
    },
]

for tc in test_cases_m2:
    add_para(f'TC ID: {tc["id"]} — {tc["title"]}')
    add_tc_table([
        ('Req Mapping', tc['req']),
        ('Priority', tc['priority']),
        ('Technique', tc['technique']),
        ('Precondition', tc['precondition']),
        ('Test Steps', tc['steps']),
        ('Test Data', tc['data']),
        ('Expected Result', tc['expected']),
        ('Pass/Fail Criteria', tc['passfail']),
    ])

# ---- M5: Workflow API ----
add_heading('MODULE M5 — WORKFLOW API (FASTAPI)', 2)

test_cases_m5 = [
    {
        'id': 'TC-M5-01',
        'title': 'Health Check Endpoint',
        'req': 'NFR-17',
        'priority': 'Medium',
        'technique': 'Functional',
        'precondition': 'Backend running on port 8000',
        'steps': '1. GET /health\n2. GET /api/health',
        'data': 'None',
        'expected': 'Both return {"status": "ok", "version": "1.0.0"} with HTTP 200',
        'passfail': 'PASS if both endpoints return 200 with status=ok',
    },
    {
        'id': 'TC-M5-02',
        'title': 'Contract Signing — Client Signs',
        'req': 'FR-7',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Contract in pending_signatures; client authenticated',
        'steps': '1. Client calls POST /contracts/{id}/sign\n2. Backend records client_signed = true\n3. Verify DB updated',
        'data': 'Contract ID = ct_e5618f617025; client = Hardhat #0',
        'expected': 'client_signed = true; status changes if both signed',
        'passfail': 'PASS if client_signed flag set in PostgreSQL',
    },
    {
        'id': 'TC-M5-03',
        'title': 'Contract Signing — Freelancer Signs',
        'req': 'FR-7',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Contract in pending_signatures; freelancer authenticated',
        'steps': '1. Freelancer calls POST /contracts/{id}/sign\n2. Backend records freelancer_signed = true\n3. Verify DB updated',
        'data': 'Freelancer = Hardhat #1',
        'expected': 'freelancer_signed = true',
        'passfail': 'PASS if freelancer_signed flag set',
    },
    {
        'id': 'TC-M5-04',
        'title': 'Contract Funding — Server-Side Transaction',
        'req': 'FR-7, NFR-14',
        'priority': 'High',
        'technique': 'Integration / Two-stage commit',
        'precondition': 'Contract in pending_funding; both signed; backend .env has CLIENT_PRIVATE_KEY',
        'steps': '1. POST /contracts/{id}/fund\n2. Backend builds and signs fundContract() tx\n3. Receipt awaited\n4. PostgreSQL updated to active',
        'data': 'Contract ID with on_chain_id set; server private key',
        'expected': 'On-chain status → InProgress; DB status → active; amounts match',
        'passfail': 'PASS if DB updated AND on-chain state changed atomically',
    },
    {
        'id': 'TC-M5-05',
        'title': 'Missing Authorization — 401 Returned',
        'req': 'FR-14',
        'priority': 'High',
        'technique': 'Negative Testing',
        'precondition': 'Backend running; no auth token',
        'steps': '1. GET /api/contracts without Authorization header\n2. Observe response',
        'data': 'No JWT token',
        'expected': 'HTTP 401; {"detail": {"code": "MISSING_AUTH_HEADER", ...}}',
        'passfail': 'PASS if 401 with error code',
    },
    {
        'id': 'TC-M5-06',
        'title': 'Milestone Approval — API Calls approveMilestone On-Chain',
        'req': 'FR-10, BR-13',
        'priority': 'High',
        'technique': 'Integration',
        'precondition': 'Milestone Submitted; client authenticated; contract has on_chain_id',
        'steps': '1. Client calls POST /milestones/{idx}/approve\n2. Backend calls approveMilestone() on contract\n3. Await receipt\n4. PostgreSQL milestone → Approved',
        'data': 'Client wallet; milestone in Submitted; contract funded',
        'expected': 'Transaction mined; PostgreSQL set to Approved; freelancer receives ETH',
        'passfail': 'PASS if full on-chain + off-chain flow succeeds',
    },
    {
        'id': 'TC-M5-07',
        'title': 'Unauthorised Milestone Approval — 403',
        'req': 'FR-10, FR-3',
        'priority': 'High',
        'technique': 'Security — RBAC',
        'precondition': 'Milestone in Submitted; freelancer (non-client) authenticated',
        'steps': '1. Freelancer calls POST /milestones/{idx}/approve\n2. Observe response',
        'data': 'Freelancer wallet token; milestone in Submitted',
        'expected': 'HTTP 403; PostgreSQL unchanged',
        'passfail': 'PASS if denied and logged',
    },
    {
        'id': 'TC-M5-08',
        'title': 'Non-Existent Contract — 404',
        'req': 'FR-12',
        'priority': 'Medium',
        'technique': 'Negative Testing',
        'precondition': 'No contract with ID = GHOST-999',
        'steps': '1. GET /api/contracts/GHOST-999\n2. Observe response',
        'data': 'Non-existent ID',
        'expected': 'HTTP 404; JSON error; no stack trace',
        'passfail': 'PASS if 404 with clean message',
    },
    {
        'id': 'TC-M5-09',
        'title': 'Proposal Accept — setFreelancer Called On-Chain',
        'req': 'FR-7, FR-16',
        'priority': 'High',
        'technique': 'Integration',
        'precondition': 'Job with on_chain_id exists; pending proposal; client authenticated',
        'steps': '1. Client calls POST /proposals/{id}/accept\n2. Backend checks contract.on_chain_id\n3. Calls setFreelancer() if exists\n4. Creates DB contract record',
        'data': 'Proposal ID; client wallet; contract.on_chain_id = 1',
        'expected': 'On-chain freelancer set; DB contract created with status = pending_signatures',
        'passfail': 'PASS if on-chain and off-chain states align',
    },
    {
        'id': 'TC-M5-10',
        'title': 'Milestone Rejection — Reason Stored',
        'req': 'FR-11',
        'priority': 'High',
        'technique': 'Functional',
        'precondition': 'Milestone in Submitted; client wallet',
        'steps': '1. Client calls POST /milestone/{id}/reject with reason\n2. Verify milestone → Rejected\n3. Check reason stored',
        'data': 'Reason = "Quality below standard"',
        'expected': 'milestone.state = Rejected; reason persisted in metadata column',
        'passfail': 'PASS if reason stored and retrievable',
    },
    {
        'id': 'TC-M5-11',
        'title': 'Reject After Approved — Blocked',
        'req': 'FR-11, FR-8',
        'priority': 'High',
        'technique': 'Negative — State boundary',
        'precondition': 'Milestone already Approved',
        'steps': '1. Client attempts to reject approved milestone\n2. Observe response',
        'data': 'Approved milestone; client token',
        'expected': 'HTTP 400; state remains Approved',
        'passfail': 'PASS if rejection blocked',
    },
    {
        'id': 'TC-M5-12',
        'title': 'PostgreSQL ↔ On-Chain State Consistency',
        'req': 'NFR-8',
        'priority': 'High',
        'technique': 'Integration',
        'precondition': 'Contract with on_chain_id; query both sources',
        'steps': '1. Query PostgreSQL for contract status\n2. Query on-chain getContractDetails\n3. Compare',
        'data': 'Any contract with on_chain_id',
        'expected': 'Both sources report identical state values',
        'passfail': 'PASS if no delta between DB and chain',
    },
]

for tc in test_cases_m5:
    add_para(f'TC ID: {tc["id"]} — {tc["title"]}')
    add_tc_table([
        ('Req Mapping', tc['req']),
        ('Priority', tc['priority']),
        ('Technique', tc['technique']),
        ('Precondition', tc['precondition']),
        ('Test Steps', tc['steps']),
        ('Test Data', tc['data']),
        ('Expected Result', tc['expected']),
        ('Pass/Fail Criteria', tc['passfail']),
    ])

# ---- SECURITY ----
add_heading('SECURITY TESTS (Cross-cutting)', 2)

test_cases_sec = [
    {
        'id': 'TC-SEC-01',
        'title': 'Third-Party Direct Contract Call — Blocked',
        'req': 'NFR-10',
        'priority': 'High',
        'technique': 'Security — Penetration',
        'precondition': 'Contract deployed; third-party wallet funded (Hardhat #2)',
        'steps': '1. Third-party calls approveMilestone() directly on-chain\n2. Observe revert',
        'data': 'Random wallet (0x3C44...); contract in Funded',
        'expected': 'Revert (onlyClient guard); no state change',
        'passfail': 'PASS if on-chain RBAC enforced independent of API',
    },
    {
        'id': 'TC-SEC-02',
        'title': 'Unauthorised Contract View — Non-Party Blocked',
        'req': 'FR-3, NFR-10',
        'priority': 'High',
        'technique': 'Security — API RBAC',
        'precondition': 'Contract exists; third-party wallet authenticated',
        'steps': '1. Third-party requests contract detail via API\n2. Observe response',
        'data': 'Wallet not in contract parties',
        'expected': 'HTTP 403; no contract metadata',
        'passfail': 'PASS if RBAC enforced',
    },
    {
        'id': 'TC-SEC-03',
        'title': 'Reentrancy in approveMilestone — Blocked',
        'req': 'FR-16, NFR-7',
        'priority': 'High',
        'technique': 'Security — CEI pattern verification',
        'precondition': 'Contract funded; milestone submitted',
        'steps': '1. Review Solidity source for CEI pattern\n2. Verify state update before external call (transfer)',
        'data': 'Source code audit of approveMilestone()',
        'expected': 'State updated before payable(freelancer).transfer(); reentrancy impossible',
        'passfail': 'PASS if CEI pattern confirmed in source',
    },
    {
        'id': 'TC-SEC-04',
        'title': 'Sensitive Information Leak Prevention',
        'req': 'NFR-11, NFR-12',
        'priority': 'Medium',
        'technique': 'Security — Inspection',
        'precondition': 'Error paths instrumented',
        'steps': '1. Trigger 401, 403, 404, 500 errors\n2. Inspect response body and server logs',
        'data': 'Invalid JWT; wrong role; non-existent resource',
        'expected': 'Generic error messages only; no stack traces; no private keys exposed',
        'passfail': 'PASS if zero sensitive leakage',
    },
]

for tc in test_cases_sec:
    add_para(f'TC ID: {tc["id"]} — {tc["title"]}')
    add_tc_table([
        ('Req Mapping', tc['req']),
        ('Priority', tc['priority']),
        ('Technique', tc['technique']),
        ('Precondition', tc['precondition']),
        ('Test Steps', tc['steps']),
        ('Test Data', tc['data']),
        ('Expected Result', tc['expected']),
        ('Pass/Fail Criteria', tc['passfail']),
    ])

# ---- SUMMARY ----
doc.add_page_break()
add_heading('TEST CASE SUMMARY', 2)
add_table = doc.add_table(rows=1, cols=2)
add_table.style = 'Light Grid Accent 1'
add_table.rows[0].cells[0].text = 'Module'
add_table.rows[0].cells[1].text = 'Test Case Count'
counts = [
    ('M1 — Wallet Auth', 2),
    ('M2 — Smart Contract Engine', 14),
    ('M5 — Workflow API', 12),
    ('Security (Cross-cutting)', 4),
    ('TOTAL', '32'),
]
for module, count in counts:
    row = add_table.add_row().cells
    row[0].text = module
    row[1].text = str(count)

output_path = '/home/sarun/Desktop/Sarun_Capstone/testing/FreeLedger_Test_Cases.docx'
doc.save(output_path)
print(f'Test Cases saved to: {output_path}')
