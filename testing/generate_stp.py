#!/usr/bin/env python3
"""Generate Software Test Plan (STP) for FreeLedger - Decentralized Freelance Protocol"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    return table

def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

# ====== TITLE PAGE ======
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Software Test Plan')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('FreeLedger: A Decentralized Freelance Protocol\nwith Web3 Integration')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Document Version: 1.0\n').bold = False
info.add_run(f'Date: {datetime.date.today().strftime("%d %B %Y")}\n')
info.add_run('Course: ITS64704 Software Testing\n')
info.add_run('Instructor: Alil Maharjan\n\n')
info.add_run('Team Members:\n').bold = True
info.add_run('Sarun Maharjan - Team Lead\n')
info.add_run('Anushree Pradhan\n')
info.add_run('Pawan Poudel\n')
info.add_run('Bijee Dangol\n')
info.add_run('Runa Maphu')

doc.add_page_break()

# ====== DOCUMENT HISTORY ======
add_heading('DOCUMENT HISTORY', 1)
add_table(
    ['Date', 'Document Version', 'Revision Description', 'Author'],
    [
        ['2026-05-10', '0.1', 'Initial STP skeleton', 'Sarun Maharjan'],
        ['2026-05-18', '0.2', 'Test cases added; scope defined', 'Anushree Pradhan'],
        ['2026-05-22', '0.3', 'Strategy and schedule completed', 'Pawan Poudel'],
        [datetime.date.today().strftime('%Y-%m-%d'), '1.0', 'Final version for submission', 'Sarun Maharjan'],
    ]
)

doc.add_page_break()

# ====== TABLE OF CONTENTS ======
add_heading('TABLE OF CONTENTS', 1)
toc_items = [
    '1. OVERVIEW',
    '   1.1 Introduction',
    '   1.2 Project Description',
    '   1.3 Testing Stakeholder Communication',
    '   1.4 Reference Material',
    '   1.5 Glossary',
    '2. TESTING CONTEXT',
    '   2.1 Test Item',
    '   2.2 Test Scope',
    '   2.2.1 Features/Functions to be Tested',
    '   2.2.2 Features/Functions NOT to be Tested',
    '3. TEST STRATEGY',
    '   3.1 Test Design Approach',
    '   3.2 Verification Approach',
    '   3.3 Validation Approach',
    '   3.3.1 Unit Test',
    '   3.3.2 Integration Test',
    '   3.3.3 System Test',
    '   3.3.4 User Acceptance Test',
    '4. TEST MANAGEMENT',
    '   4.1 Administration',
    '   4.2 Approval Authority',
    '5. PROJECT SCHEDULE AND TEST ARTIFACT REPOSITORY',
    '   5.1 Test Schedule',
    '   5.2 RACI Matrix',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ====== 1. OVERVIEW ======
add_heading('1. OVERVIEW', 1)

add_heading('1.1 Introduction', 2)
add_para(
    'This Software Test Plan (STP) documents the testing strategy, scope, resources, and schedule '
    'for the FreeLedger platform — a decentralized freelance protocol with Web3 integration. '
    'The system uses Ethereum smart contracts (Solidity), a FastAPI backend, React frontend, '
    'and PostgreSQL to enable trustless escrow-based freelance engagements.\n\n'
    'This STP focuses specifically on the blockchain contract lifecycle module (Module M2) and '
    'the workflow API layer (Module M5) that together govern the core value proposition: '
    'creating, funding, executing, and resolving smart-contract escrows for milestone-based freelance work.'
)

add_heading('1.2 Project Description', 2)
add_para(
    'FreeLedger is a decentralized freelance platform that eliminates intermediary trust by using '
    'Ethereum smart contracts as binding escrow agreements. Key capabilities include:\n\n'
    '• Wallet-based authentication via MetaMask (EIP-712 signatures)\n'
    '• On-chain escrow contract deployment per job\n'
    '• Milestone-based payment release with 2.5% platform fee\n'
    '• Role-based access for clients, freelancers, and admins\n'
    '• Dispute initiation and resolution on-chain\n'
    '• PostgreSQL metadata indexing for fast query\n\n'
    'The system architecture consists of:\n'
    '  - Smart Contract Layer: GigEscrow.sol (Solidity, deployed on Hardhat local chain)\n'
    '  - Backend API Layer: FastAPI (Python) with Web3.py integration\n'
    '  - Frontend Layer: React with ethers.js and MetaMask\n'
    '  - Data Layer: PostgreSQL + Redis for session management\n\n'
    'Business Objectives:\n'
    '  - Reduce freelance payment disputes through immutable on-chain escrow\n'
    '  - Enable global, permissionless freelancer-client engagements\n'
    '  - Achieve sub-60-second full state change lifecycle\n'
    '  - Maintain 99% API availability'
)

add_heading('1.3 Testing Stakeholder Communication', 2)
add_para(
    'Testing communication follows a structured escalation framework:\n\n'
    'Weekly Status Meetings: Every Monday at 10:00 AM, the test team presents status, '
    'open defects, and blocking issues to the project manager.\n\n'
    'Defect Triage: Daily 15-minute standup to prioritize newly discovered defects.\n\n'
    'Escalation Path:\n'
    '  1. Tester → Test Lead: Issue resolution within 24 hours\n'
    '  2. Test Lead → Test Manager: If unresolved, escalated within 48 hours\n'
    '  3. Test Manager → Project Manager: Final arbitration for critical blockers\n\n'
    'Communication Channels:\n'
    '  - Discord: Daily updates and defect discussion\n'
    '  - GitHub Issues: Formal defect tracking\n'
    '  - Email: Weekly status report distribution\n'
    '  - Shared Google Drive: Test artifact repository'
)

add_heading('1.4 Reference Material', 2)
add_table(
    ['Document', 'Source'],
    [
        ['0371759_Sarun_Capstone.pdf', 'Primary SRS / Capstone Report'],
        ['GigEscrow.sol', 'Smart contract source (317 lines)'],
        ['backend/app/routers/contracts.py', 'Contract lifecycle API endpoints'],
        ['backend/app/services/blockchain_service.py', 'Blockchain interaction layer'],
        ['backend/app/schemas.py', 'Pydantic data models'],
        ['IEEE 829-2008', 'Standard for Software and System Test Documentation'],
        ['ISO/IEC/IEEE 29119-1/2/3', 'Software Testing Standards'],
        ['ISTQB Glossary v2.3', 'Standard testing terminology'],
    ]
)

add_heading('1.5 Glossary', 2)
add_table(
    ['Term', 'Definition'],
    [
        ['Escrow', 'Smart contract that holds funds until conditions met'],
        ['Milestone', 'A discrete deliverable with associated payment amount'],
        ['On-chain ID', 'Integer ID assigned by the smart contract to each escrow'],
        ['setFreelancer', 'Contract function that assigns the freelancer to an escrow'],
        ['approveMilestone', 'Contract function releasing payment to freelancer (minus fee)'],
        ['Hardhat', 'Local Ethereum development network'],
        ['MetaMask', 'Browser wallet for Ethereum transaction signing'],
        ['EIP-712', 'Typed structured data signing standard'],
        ['CEI', 'Checks-Effects-Interactions pattern (reentrancy guard)'],
    ]
)

doc.add_page_break()

# ====== 2. TESTING CONTEXT ======
add_heading('2. TESTING CONTEXT', 1)

add_heading('2.1 Test Item', 2)
add_para(
    'The primary test item is the blockchain contract lifecycle module of FreeLedger, '
    'encompassing:\n\n'
    '1. GigEscrow.sol smart contract (deployed at 0x5FbDB2315678afecb367f032d93F642f64180aa3)\n'
    '2. blockchain_service.py — transaction building, signing, and broadcasting\n'
    '3. contract_service.py — sign/fund orchestration with two-stage commit\n'
    '4. contracts.py router — REST endpoints for contract CRUD and state transitions\n'
    '5. proposals.py router — proposal accept handler with on-chain freelancer assignment\n'
    '6. PostgreSQL contract_milestones and contracts tables\n\n'
    'Pass/Fail Criteria:\n'
    '  - All High-priority test cases: 100% pass rate required\n'
    '  - Medium-priority test cases: ≥ 90% pass rate\n'
    '  - Zero Critical (Severity 1) open defects at exit\n'
    '  - No regression in previously passing functionality'
)

add_heading('2.2 Test Scope', 2)

add_heading('2.2.1 Features/Functions to be Tested', 3)
add_para(
    'In-Scope features:\n\n'
    'Smart Contract (M2):\n'
    '  - Contract deployment and existence check\n'
    '  - setFreelancer() access control (onlyClient modifier)\n'
    '  - fundContract() with correct value and caller\n'
    '  - submitMilestone() by assigned freelancer\n'
    '  - approveMilestone() with automatic fee deduction\n'
    '  - Reentrancy prevention (CEI pattern)\n'
    '  - State machine transitions: Created → InProgress → Completed\n'
    '  - Double-approve prevention\n'
    '  - Invalid caller rejection (non-client, non-freelancer)\n\n'
    'Workflow API (M5):\n'
    '  - Contract creation with on_chain_id\n'
    '  - Contract signing (client and freelancer)\n'
    '  - Contract funding (server-side transaction)\n'
    '  - Milestone state transitions (Pending → Submitted → Approved)\n'
    '  - Milestone rejection with reason storage\n'
    '  - Health check endpoints\n\n'
    'Cross-cutting:\n'
    '  - RBAC: role-based access to contract operations\n'
    '  - PostgreSQL/Blockchain state consistency\n'
    '  - Authentication via MetaMask (M1)\n'
    '  - Frontend contract detail display (M7)'
)

add_heading('2.2.2 Features/Functions NOT to be Tested', 3)
add_para(
    'Out-of-Scope for this test cycle:\n\n'
    '  - IPFS storage module (M3) — CID integration not yet live\n'
    '  - Matching engine (M6) — rule-based freelancer matching\n'
    '  - Admin dashboard (separate sprint)\n'
    '  - Load/performance testing beyond 10 concurrent requests\n'
    '  - Cross-browser compatibility (Chrome only for this cycle)\n'
    '  - Mobile responsiveness\n'
    '  - Production-scale PostgreSQL indexing (10M+ rows)'
)

doc.add_page_break()

# ====== 3. TEST STRATEGY ======
add_heading('3. TEST STRATEGY', 1)

add_heading('3.1 Test Design Approach', 2)
add_para(
    'Test design follows a requirement-driven approach with traceability to the 38 requirements '
    'defined in the RTM (Requirements Traceability Matrix). Techniques employed:\n\n'
    '• Equivalence Partitioning (EP): For input validation, role-based access\n'
    '• Boundary Value Analysis (BVA): For milestone amounts, gas limits, array indices\n'
    '• Decision Table (DT): For state transition combinations and dispute outcomes\n'
    '• White Box: Branch coverage for Solidity modifiers (onlyClient, onlyFreelancer, require())\n'
    '• Negative Testing: Invalid callers, zero values, non-existent IDs\n'
    '• Security Testing: Reentrancy, replay attacks, unauthorized access\n\n'
    'Test cases are organized by module (M1-M8) with cross-cutting Security and Consistency suites.'
)

add_table(
    ['Entry Criteria', 'Description'],
    [
        ['Requirements', 'All requirements documented and approved'],
        ['Environment', 'Hardhat node running, PostgreSQL seeded, backend deployed'],
        ['Test Data', 'Test wallets funded (Hardhat accounts 0-9)'],
        ['Test Cases', 'Test cases written, reviewed, and approved'],
        ['Access', 'All team members have environment access'],
    ]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [
        ['Pass Rate', 'High-priority: 100%, Medium: >= 90%'],
        ['Defects', 'No Critical (S1) or Major (S2) open defects'],
        ['Coverage', '100% requirement coverage in RTM'],
        ['Documentation', 'Test execution report submitted and approved'],
    ]
)
add_para('')
add_table(
    ['Deliverable', 'Due Date'],
    [
        ['Test Plan Document', '2026-05-15'],
        ['Test Cases (RTM)', '2026-05-18'],
        ['Test Execution Report', '2026-05-22'],
        ['Defect Log', '2026-05-22'],
    ]
)

add_heading('3.2 Verification Approach', 2)

add_heading('3.2.1 Requirements Review', 3)
add_para(
    'Requirements were reviewed against:\n'
    '  - Capstone report (0371759_Sarun_Capstone.pdf)\n'
    '  - Solidity smart contract source\n'
    '  - FastAPI router implementations\n\n'
    'Each requirement was verified to be testable — i.e., has clear pass/fail criteria, '
    'observable outputs, and controllable inputs.'
)
add_table(
    ['Entry Criteria', 'Description'],
    [
        ['Requirements documented', 'Per IT standards template'],
        ['SRS approved', 'Signed off by project sponsor'],
    ]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [
        ['Requirements testable', 'All 38 requirements mapped to >= 1 TC'],
        ['Traceability matrix', 'RTM complete with TC-ID mapping'],
    ]
)
add_para('')
add_table(
    ['Deliverable', 'Due Date'],
    [['RTM Document', '2026-05-10']]
)

add_heading('3.2.2 Use Case Review', 3)
add_para(
    'Use cases were reviewed for the contract lifecycle:\n'
    '  1. Client posts job → MetaMask deploys escrow → on-chain ID stored\n'
    '  2. Freelancer proposes → client accepts → setFreelancer called\n'
    '  3. Both parties sign → client funds → contract active\n'
    '  4. Freelancer submits milestone → client approves → ETH released\n'
    '  5. Dispute raised → admin resolves → funds distributed'
)
add_table(
    ['Entry Criteria', 'Description'],
    [['Use cases documented', 'Per project template, reviewed by team']]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [['Use cases approved', 'Signed off by all team members']]
)

add_heading('3.2.3 Design Review', 3)
add_para(
    'Design review focused on:\n'
    '  - Smart contract state machine correctness (enum ContractStatus)\n'
    '  - Two-stage commit pattern (blockchain tx → PostgreSQL update)\n'
    '  - Reentrancy protection via CEI pattern\n'
    '  - Access control modifiers (onlyClient, onlyFreelancer)'
)
add_table(
    ['Entry Criteria', 'Description'],
    [['Design documented', 'Solidity source + Flow diagrams in capstone report']]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [['Design approved', 'Reviewed by team, aligns with requirements']]
)

add_heading('3.3 Validation Approach', 2)

add_heading('3.3.1 Unit Test', 3)
add_para(
    'Unit testing covers the smallest testable components:\n\n'
    'Smart Contract (via Hardhat tests):\n'
    '  - Individual function calls with valid/invalid inputs\n'
    '  - Modifier branch coverage (onlyClient, onlyFreelancer)\n'
    '  - require() revert messages\n'
    '  - State transition guard conditions\n\n'
    'Backend (via pytest):\n'
    '  - Pydantic schema validation\n'
    '  - Service function edge cases\n'
    '  - Error handling and exception paths'
)
add_table(
    ['Entry Criteria', 'Description'],
    [
        ['Source code', 'Coded following project standards'],
        ['Unit tests', 'Written for all smart contract functions'],
    ]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [
        ['Tests passing', 'All unit tests pass'],
        ['No S1/S2 defects', 'Zero critical or major open defects'],
    ]
)

add_heading('3.3.2 Integration Test', 3)
add_para(
    'Integration testing validates the interaction between:\n\n'
    '• Backend ↔ Smart Contract: blockchain_service.py makes eth_call and send_transaction\n'
    '• Backend ↔ PostgreSQL: contract_service.py performs two-stage commit\n'
    '• Frontend ↔ Backend: REST API calls from React to FastAPI\n'
    '• MetaMask ↔ Frontend: EIP-712 signing and transaction confirmation\n\n'
    'Each integration test exercises the full boundary between two layers, verifying '
    'data format, error propagation, and state consistency.'
)
add_table(
    ['Entry Criteria', 'Description'],
    [
        ['Unit tests passed', 'Evidence that unit test goals were completed'],
        ['Environment', 'All services (Hardhat, PostgreSQL, Backend) running'],
    ]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [
        ['S1/S2 defects closed', 'All critical and major integration defects resolved'],
        ['Pass rate', '100% on high-priority integration test cases'],
    ]
)

add_heading('3.3.3 System Test', 3)
add_para(
    'System testing validates end-to-end scenarios across all layers:\n\n'
    'Test Categories:\n'
    '  - Security: RBAC, reentrancy, unauthorized access logging\n'
    '  - Functional: Full contract lifecycle from job posting to milestone approval\n'
    '  - Consistency: PostgreSQL ↔ on-chain state synchronization\n'
    '  - Negative: Invalid inputs, wrong caller, non-existent resources\n'
    '  - Boundary: Zero values, maximum milestone amounts, address(0)'
)
add_table(
    ['Category', 'Required', 'Test Objective'],
    [
        ['Functional', 'Yes', 'Verify complete contract lifecycle'],
        ['Security', 'Yes', 'Verify RBAC and reentrancy guards'],
        ['Consistency', 'Yes', 'Verify DB matches blockchain state'],
        ['Negative', 'Yes', 'Verify graceful error handling'],
        ['Performance', 'No', 'Deferred to next cycle'],
        ['Load', 'No', 'Deferred to next cycle'],
    ]
)
add_para('')
add_table(
    ['Entry Criteria', 'Description'],
    [
        ['Integration tests passed', 'Evidence of successful integration test completion'],
        ['Environment stable', 'All services healthy for 24+ hours'],
    ]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [
        ['S1/S2 defects closed', 'All critical and major system test defects resolved'],
        ['Pass rate', 'High-priority: 100%, Medium: >= 90%'],
    ]
)

add_heading('3.3.4 User Acceptance Test (UAT)', 3)
add_para(
    'UAT will be performed by simulating real user roles:\n'
    '  - Client: Posts a job, funds escrow, approves milestones\n'
    '  - Freelancer: Proposes, submits work, receives payment\n'
    '  - Admin: Monitors disputes, resolves conflicts\n\n'
    'UAT scenarios follow the defined use cases and require manual MetaMask interaction.'
)
add_table(
    ['Entry Criteria', 'Description'],
    [
        ['System tests passed', 'All system test exit criteria met'],
        ['Testers trained', 'All testers familiar with MetaMask and system flow'],
    ]
)
add_para('')
add_table(
    ['Exit Criteria', 'Description'],
    [
        ['UAT defects closed', 'All S1/S2 defects from UAT resolved'],
        ['Sign-off', 'Project sponsor approves UAT completion'],
    ]
)

doc.add_page_break()

# ====== 4. TEST MANAGEMENT ======
add_heading('4. TEST MANAGEMENT', 1)

add_heading('4.1 Administration', 2)
add_para(
    'Test management is coordinated by the Test Lead (Sarun Maharjan).\n\n'
    'Tools:\n'
    '  - GitHub Issues: Defect tracking\n'
    '  - This STP document: Test planning\n'
    '  - Manual execution logs: Test execution tracking\n'
    '  - Discord: Team communication\n\n'
    'Defect Severity Classification:\n'
    '  - Critical (S1): Contract funds at risk, data loss, security breach\n'
    '  - Major (S2): Feature broken, no workaround\n'
    '  - Minor (S3): Feature works with limitations, cosmetic issue\n'
    '  - Trivial (S4): Cosmetic, documentation, non-functional'
)

add_heading('4.2 Approval Authority', 2)
add_table(
    ['Test Activity', 'Authorized Approver', 'Reference'],
    [
        ['Test Design Approach', 'Sarun Maharjan (Test Lead)', 'Section 3.1'],
        ['Requirements Review', 'Sarun Maharjan', 'Section 3.2.1'],
        ['Use Case Review', 'Anushree Pradhan', 'Section 3.2.2'],
        ['Design Review', 'Pawan Poudel', 'Section 3.2.3'],
        ['Unit Test', 'Pawan Poudel', 'Section 3.3.1'],
        ['Integration Test', 'Bijee Dangol', 'Section 3.3.2'],
        ['System Test', 'Runa Maphu', 'Section 3.3.3'],
        ['UAT', 'Sarun Maharjan', 'Section 3.3.4'],
        ['Test Summary Report', 'Sarun Maharjan', 'Execution Report'],
    ]
)

doc.add_page_break()

# ====== 5. PROJECT SCHEDULE ======
add_heading('5. PROJECT SCHEDULE AND TEST ARTIFACT REPOSITORY', 1)

add_heading('5.1 Test Schedule', 2)
add_table(
    ['Milestone', 'Start Date', 'Completion Date'],
    [
        ['Test Plan Complete', '2026-05-01', '2026-05-10'],
        ['Test Case Design (RTM)', '2026-05-10', '2026-05-18'],
        ['Unit Test Execution', '2026-05-15', '2026-05-18'],
        ['Integration Test Execution', '2026-05-18', '2026-05-20'],
        ['System Test Execution', '2026-05-20', '2026-05-22'],
        ['UAT Execution', '2026-05-22', '2026-05-25'],
        ['Defect Triage & Retest', '2026-05-22', '2026-05-28'],
        ['Test Summary Report', '2026-05-28', '2026-05-30'],
    ]
)

add_heading('5.2 RACI Matrix', 2)
add_table(
    ['Activity', 'Sarun', 'Anushree', 'Pawan', 'Bijee', 'Runa'],
    [
        ['Test Plan', 'A/R', 'C', 'C', 'I', 'I'],
        ['Test Case Design', 'A', 'R', 'R', 'C', 'C'],
        ['Unit Testing (Contract)', 'A', 'I', 'R', 'C', 'C'],
        ['Integration Testing', 'A', 'R', 'C', 'R', 'C'],
        ['System Testing', 'A', 'C', 'C', 'C', 'R'],
        ['Defect Management', 'A', 'R', 'R', 'R', 'R'],
        ['Execution Report', 'A/R', 'C', 'C', 'C', 'C'],
    ]
)
add_para('R = Responsible, A = Accountable, C = Consulted, I = Informed')

doc.add_page_break()

# ====== DOCUMENT APPROVAL ======
add_heading('DOCUMENT APPROVAL HISTORY', 1)
add_table(
    ['Role', 'Name', 'Signature', 'Date'],
    [
        ['Prepared By', 'Sarun Maharjan', '______', datetime.date.today().strftime('%Y-%m-%d')],
        ['Reviewed By', 'Anushree Pradhan', '______', '______'],
        ['Reviewed By', 'Pawan Poudel', '______', '______'],
        ['Approved By', 'Alil Maharjan (Instructor)', '______', '______'],
    ]
)

output_path = '/home/sarun/Desktop/Sarun_Capstone/testing/FreeLedger_Test_Plan.docx'
doc.save(output_path)
print(f'STP saved to: {output_path}')
