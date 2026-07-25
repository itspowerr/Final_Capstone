from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_OUT = ROOT / "Capstone-II_Frontend_Bijee_Dangol_Revised.docx"
DOWNLOADS_OUT = Path(r"C:\Users\bishe\Downloads\Capstone-II_Frontend_Bijee_Dangol_Revised.docx")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for para in cell.paragraphs:
        para.paragraph_format.space_after = Pt(2)


def make_table(doc, headers, rows, font_size=8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, header in enumerate(headers):
        cell_text(table.rows[0].cells[i], header, bold=True, size=font_size)
        shade(table.rows[0].cells[i])
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=font_size)
    doc.add_paragraph("")
    return table


def para(doc, text="", bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("â€¢ " + text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F2F2F2")
    p._p.get_or_add_pPr().append(shd)


def add_img(doc, rel_path, caption):
    para(
        doc,
        f"Evidence reference: {caption}. The screenshot file is stored separately at {rel_path}; it is not enlarged inside the report to keep the document readable.",
    )


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INDIVIDUAL CAPSTONE-II REPORT")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Frontend Implementation and Testing of FreeLedger")
    r.bold = True
    r.font.size = Pt(15)
    r.font.name = "Times New Roman"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project: Decentralized Freelance Protocol with Web3 Integration")
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"

    make_table(
        doc,
        ["Field", "Details"],
        [
            ["Student Name", "Bijee Dangol"],
            ["Student ID", "0372053"],
            ["Role", "Frontend Developer"],
            [
                "Report Scope",
                "Frontend implementation, frontend-backend integration, browser testing, and frontend critical evaluation only",
            ],
            ["Prepared Date", "24 July 2026"],
        ],
        font_size=10,
    )
    para(
        doc,
        "This Capstone-II individual report continues from the first capstone report. The first report focused mainly on literature review, analysis, requirements, frontend architecture, and early UI direction. This second report documents the actual frontend implementation work, the changes made during development, the problems faced while integrating and testing the application, and the final frontend testing evidence. Backend, database, smart-contract, and infrastructure details are discussed only where they directly affected frontend integration and user experience.",
    )
    doc.add_page_break()


def chapter_5(doc):
    doc.add_heading("Chapter 5: Implementation", 1)
    doc.add_heading("5.1 Overview", 2)
    para(
        doc,
        "The implementation phase converted the planned frontend design of FreeLedger into a working React application. My individual responsibility was the frontend layer: public landing page, authentication interface, role-based dashboards, navigation, job browsing, proposal submission interface, contracts pages, profile pages, messages page, admin-facing pages, frontend API integration, user feedback states, responsive layout, and automated frontend testing support.",
    )
    para(
        doc,
        "During Capstone-I, the frontend was described mainly as a React-based interface that would simplify blockchain and freelance-platform interactions. In Capstone-II, that design changed from a static plan into a functional multi-page application. The frontend now communicates with the FastAPI backend through Axios, stores authentication/session data in browser storage, redirects users according to role, renders separate client and freelancer workspaces, and includes frontend tests and browser automation evidence.",
    )
    para(
        doc,
        "The objective was not only to make screens visually complete, but also to make the frontend usable as a working system. This required solving integration issues such as API URL configuration, CORS errors, local development setup differences, authentication flow changes, validation behavior, React routing behavior, and testing tool compatibility.",
    )

    doc.add_heading("5.2 Development Environment", 2)
    doc.add_heading("5.2.1 Hardware", 3)
    make_table(
        doc,
        ["Hardware Item", "Specification / Purpose"],
        [
            ["Laptop / Development Machine", "Windows-based local development system used to run React frontend, FastAPI backend, Docker services, and browser tests."],
            ["Processor and Memory", "General development hardware sufficient for running Node.js, Python backend, Docker containers, and Chromium browser automation simultaneously."],
            ["Storage", "Local disk storage used for source code, node_modules, Python virtual environment, screenshots, Word report files, and test evidence."],
            ["Network", "Localhost development network using frontend at port 3000, backend at port 8000, Hardhat at port 8545, and IPFS gateway at port 8080."],
        ],
    )
    doc.add_heading("5.2.2 Software", 3)
    make_table(
        doc,
        ["Software / Tool", "Use in Frontend Implementation"],
        [
            ["React.js", "Main frontend framework used to build reusable components and route-based pages."],
            ["JavaScript, HTML, CSS", "Used for UI logic, markup, styling, responsive layouts, and component behavior."],
            ["React Router", "Used for routing between landing page, login page, client dashboard, freelancer dashboard, profile, contracts, jobs, messages, and admin pages."],
            ["Axios", "Used as the frontend HTTP client for backend API requests. Axios interceptors attach authentication tokens and handle token refresh/logout behavior."],
            ["Google Identity Services", "Added to support Google login/signup button rendering when a Google client ID is configured."],
            ["Browser Local Storage", "Used to store access token, refresh token, and user role data for session persistence and role-based redirection."],
            ["Jest and React Testing Library", "Used to create and run unit/component tests for landing navigation and login form validation/authentication behavior."],
            ["Playwright Chromium", "Used to run real browser smoke tests and end-to-end freelancer workflow tests."],
            ["Docker Services", "PostgreSQL, Redis, IPFS, and Hardhat were required locally so the frontend could be tested against the real backend stack."],
            ["Git and GitHub", "Used for version control and pushing completed frontend/testing changes to the project repository."],
        ],
    )

    doc.add_heading("5.3 System Architecture", 2)
    para(
        doc,
        "The frontend architecture follows a layered React structure. The route layer maps browser paths to page components. The page layer renders complete screens for users. The component layer contains reusable navigation bars, modals, cards, notifications, and shared UI sections. The service layer contains Axios and configuration logic. This separation helped keep the UI maintainable while allowing API and session behavior to be reused across the application.",
    )
    make_table(
        doc,
        ["Layer", "Frontend Responsibility", "Implemented Examples"],
        [
            ["Routing Layer", "Maps browser routes to page components and separates normal user mode from admin mode.", "App.js; /login; /client/dashboard; /freelancer/jobs; admin routes."],
            ["Page Layer", "Renders complete screens for each user workflow.", "Landing.js, Login.js, Client Dashboard, Freelancer Dashboard, Messages, Profiles, Contracts, Jobs."],
            ["Component Layer", "Provides reusable UI sections such as navbars, modal forms, dashboard cards, notifications, and user menus.", "Client Navbar, Freelancer Navbar, Admin Navbar, PostProjectModal."],
            ["State and Context Layer", "Stores shared frontend state such as wallet status and app-level user/session behavior.", "AppContext.js and useAuth hook."],
            ["API Service Layer", "Controls backend base URL, token injection, refresh behavior, and logout on failed refresh.", "services/api.js and config/index.js."],
            ["Styling Layer", "Provides visual identity, responsive layout, role-specific dashboard styling, and transitions.", "landing.css, login.css, client CSS, freelancer CSS, admin CSS, motion.css."],
        ],
    )
    make_table(
        doc,
        ["Area", "Confirmed Frontend Routes"],
        [
            ["Public", "/ and /login"],
            ["Client", "/client/dashboard, /client/explore-jobs, /client/browse-freelancers, /client/my-contracts, /client/profile, /client/messages"],
            ["Freelancer", "/freelancer/dashboard, /freelancer/jobs, /freelancer/contracts, /freelancer/my-profile, /freelancer/messages"],
            ["Admin Mode", "/login, /dashboard, /users, /jobs, /proposals, /contracts, /disputes, /audit-logs"],
        ],
    )

    doc.add_heading("5.4 Implementation Details", 2)
    doc.add_heading("5.4.1 Modules Implemented", 3)
    make_table(
        doc,
        ["Module", "Frontend Implementation Details", "Purpose / User Value"],
        [
            ["Landing Page", "Redesigned public homepage with hero content, proof points, feature sections, safety messaging, calls to action, and navigation to authentication.", "Introduces FreeLedger clearly and gives users a less technical first impression of a Web3 freelancing platform."],
            ["Authentication Page", "Tab-based Sign In / Create Account screen with email/password login, registration, role selection, loading states, error popups, and Google button rendering.", "Allows clients and freelancers to enter the platform and reach the correct dashboard."],
            ["Session Persistence", "Successful authentication stores access token, refresh token, and user information in localStorage and redirects users based on role.", "Keeps the user session available across frontend pages and enables authenticated API calls."],
            ["Axios API Service", "Central Axios instance with REACT_APP_API_URL fallback, JSON headers, bearer token injection, refresh handling, and logout behavior.", "Prevents duplicated API logic and connects the frontend consistently with the backend."],
            ["Client Dashboard", "Client workspace with statistics, active project/contract display, applied freelancer summary, export report, notifications, profile menu, and navigation.", "Gives clients a central workspace for managing projects, proposals, contracts, and profile actions."],
            ["Freelancer Dashboard", "Role-specific navigation, dashboard view, job browsing, contracts, profile, messages, notifications, and user menu.", "Gives freelancers a separate workspace focused on finding work and managing profile/contracts."],
            ["Find Jobs and Proposal UI", "Search/filter controls, job cards, job details modal, cover letter input, bid amount input, estimated-days input, and proposal submission.", "Allows freelancers to browse jobs and submit proposals through a real frontend workflow."],
            ["Messages Page", "Shared messages route for client and freelancer message access.", "Supports the communication requirement of the freelancing platform."],
            ["Profile Pages", "Separate client and freelancer profile pages.", "Allows users to manage personal or professional information according to role."],
            ["Admin Frontend Pages", "Admin login, dashboard, users, jobs, proposals, contracts, disputes, and audit logs pages in admin mode.", "Provides an administrative frontend for monitoring and management."],
            ["Automated Frontend Tests", "Jest tests and Playwright scripts for login validation, navigation, browser smoke testing, and E2E freelancer proposal workflow.", "Makes frontend behavior verifiable instead of relying only on manual checking."],
        ],
        font_size=7,
    )

    doc.add_heading("5.4.2 Authentication Changes", 3)
    para(
        doc,
        "One of the biggest changes from the earlier design was the authentication flow. Initially, the project was strongly focused on wallet-based authentication. During implementation, the frontend was expanded so users could also register and log in using traditional email/password authentication, and Google authentication was added as an optional method. This improved usability because normal users can access the platform without first understanding wallet setup.",
    )
    para(
        doc,
        "The login form validates email and password before calling the backend. Registration checks first name, email, password length of at least eight characters, and role selection. After successful login or registration, the frontend stores tokens and user details, then routes the user to either the client or freelancer dashboard.",
    )
    code(
        doc,
        "const persistSession = useCallback(({ access_token, refresh_token, user }) => {\n"
        "  localStorage.setItem('access_token', access_token);\n"
        "  localStorage.setItem('refresh_token', refresh_token);\n"
        "  localStorage.setItem('user', JSON.stringify(user));\n"
        "  navigate(user.role === 'client' ? '/client/dashboard' : '/freelancer/dashboard');\n"
        "}, [navigate]);",
    )
    para(
        doc,
        "Google login was implemented by loading the Google Identity Services script only when a Google client ID is configured. When a Google credential is received, the frontend sends it to /auth/google together with the selected role.",
    )

    doc.add_heading("5.4.3 API Integration", 3)
    para(
        doc,
        "The frontend-backend connection is handled through a central Axios service. The frontend uses REACT_APP_API_URL when provided and falls back to http://127.0.0.1:8000/api for local development. This makes local use possible while still allowing deployed projects to set their own backend URL.",
    )
    code(
        doc,
        "const api = axios.create({\n"
        "  baseURL: process.env.REACT_APP_API_URL || 'http://127.0.0.1:8000/api',\n"
        "  headers: { 'Content-Type': 'application/json' },\n"
        "});\n\n"
        "api.interceptors.request.use((config) => {\n"
        "  const token = localStorage.getItem('access_token');\n"
        "  if (token) config.headers.Authorization = `Bearer ${token}`;\n"
        "  return config;\n"
        "});",
    )
    para(
        doc,
        "The API service also supports refresh-token behavior. When an authenticated request returns 401, the frontend attempts to refresh the token. If refresh fails, local session data is removed and the user is redirected to /login.",
    )

    doc.add_heading("5.4.4 Role-Based UI and Navigation", 3)
    para(
        doc,
        "The implemented frontend separates users by role. Clients and freelancers do not see the same dashboard routes. Clients are routed to /client/dashboard, while freelancers are routed to /freelancer/dashboard. Separate client and freelancer navigation bars reduce confusion because each role sees actions related to their own workflow.",
    )

    doc.add_heading("5.4.5 Visual and Responsive Implementation", 3)
    para(
        doc,
        "Compared with the initial report, the frontend changed from basic planned pages into a more complete visual product. Separate CSS files were created for landing, login, messages, client pages, freelancer pages, admin pages, navigation, and motion effects. Responsive browser checks were also added for desktop, tablet, and mobile viewports.",
    )
    add_img(doc, "testing/browser-screenshots/homepage-desktop.png", "Figure 5.1: FreeLedger homepage rendered in desktop browser test")
    add_img(doc, "testing/browser-screenshots/homepage-mobile.png", "Figure 5.2: FreeLedger homepage rendered in mobile browser test")

    doc.add_heading("5.4.6 Changes from Capstone-I to Capstone-II", 3)
    make_table(
        doc,
        ["Area", "Capstone-I Direction", "Capstone-II Implemented Change"],
        [
            ["Frontend Scope", "React architecture, login page, landing page, and dashboard concepts.", "Expanded into client, freelancer, shared message, profile, contract, job, proposal, and admin interfaces."],
            ["Authentication", "Wallet/Web3-oriented login concept.", "Email/password, registration, role-based login, and Google authentication support were added."],
            ["User Flow", "Expected interaction diagrams and page-level component planning.", "Working role redirects, dashboard navigation, job browsing, proposal modal, profile access, messages, and logout."],
            ["Backend Connection", "Service adapters were described conceptually.", "Central Axios API service with environment URL, token injection, refresh handling, and CORS/API URL troubleshooting."],
            ["Testing", "Mostly analysis and design stage.", "Jest unit tests, integration checks, Playwright browser tests, screenshots, and E2E freelancer workflow evidence."],
            ["Problems Discovered", "UI complexity and Web3 usability were identified as risks.", "Actual implementation exposed CORS, API URL, route fallback, testing framework, local environment, and accessibility issues."],
        ],
        font_size=7,
    )

    doc.add_heading("5.5 Algorithms / Techniques Used", 2)
    make_table(
        doc,
        ["Technique", "How It Was Used in the Frontend", "Reason for Use"],
        [
            ["Component-Based Architecture", "Pages and UI sections were separated into React components.", "Improves maintainability and avoids one large frontend file."],
            ["State Management with Hooks", "useState, useEffect, useCallback, useMemo, and useRef handled form values, loading states, modals, notifications, and API data.", "Allows dynamic UI updates without reloading the page."],
            ["Role-Based Routing", "The frontend checks authenticated user role and redirects to client or freelancer dashboards.", "Ensures users reach the correct workspace."],
            ["Centralized API Service", "Axios handles base URL, JSON headers, Authorization token, refresh token, and logout behavior.", "Keeps backend communication consistent."],
            ["Conditional Rendering", "Google button, validation errors, loading labels, modals, notifications, and dashboard content render based on state.", "Improves user feedback and prevents unavailable features from appearing active."],
            ["Form Validation", "Login and registration forms validate fields before API requests.", "Reduces unnecessary backend calls and gives immediate feedback."],
            ["Responsive Design", "CSS layout and browser testing verified desktop, tablet, and mobile rendering.", "Improves usability across screen sizes."],
            ["Browser Automation", "Playwright scripts performed real browser login, dashboard navigation, and freelancer proposal submission.", "Provides stronger evidence than manual observation alone."],
        ],
        font_size=7,
    )

    doc.add_heading("5.6 Challenges Faced and How They Were Resolved", 2)
    para(
        doc,
        "Several technical problems were faced during frontend implementation and testing. These problems were important because they affected real usability. A user may see only a simple error such as Network Error, but the underlying cause can be API URL mismatch, CORS configuration, backend not running, or browser security restrictions.",
    )
    make_table(
        doc,
        ["Problem Faced", "Frontend Impact", "Resolution / Current Status"],
        [
            ["Login showed Network Error during local use.", "The user could not register or log in because the browser could not complete the API call.", "The frontend API URL was checked and set to http://127.0.0.1:8000/api. Backend health and CORS preflight were verified."],
            ["CORS preflight returned OPTIONS 400 for /api/auth/register.", "The browser blocked the register request before the actual POST request was sent.", "Backend CORS origins were configured to allow http://localhost:3000 and http://127.0.0.1:3000."],
            ["Confusion between local and deployed API URLs.", "A cloned/deployed frontend can fall back to localhost/127.0.0.1, which points to the current machine.", "Setup guidance now clarifies local API URL versus deployed backend URL. .env files are intentionally not pushed."],
            ["Python 3.14 virtual environment and package DLL blocking affected backend startup during integration testing.", "The frontend could not be tested if the backend failed to run locally.", "The solution was to recreate the backend virtual environment using Python 3.11/3.12 and address Windows security blocking when necessary."],
            ["Google authentication required external Google sign-in for full verification.", "Google button and invalid-token handling could be tested, but real OAuth success required human consent.", "Google button rendering was tested. Invalid token handling was verified. Real OAuth was documented as blocked without test Google sign-in."],
            ["Jest did not discover tests through the standard CRA command in the hidden .gemini Windows path.", "The test files existed, but CRA/Jest returned No tests found.", "A custom Jest config and clean-path workaround allowed all Jest tests to run successfully."],
            ["React Router v7 needed compatibility fixes in Jest.", "Tests initially failed because Jest/CRA could not resolve React Router modules and lacked TextEncoder.", "A Jest config mapping and testPolyfills.js were added."],
            ["Direct HTTP access to React routes returned 404 in command-line checks.", "Routes such as /login and /client/dashboard returned 404 when requested directly, although browser navigation worked.", "The defect was preserved for future fix. A route fallback should be configured for production/static serving."],
            ["Login form labels were not programmatically linked to inputs.", "React Testing Library label queries exposed an accessibility weakness.", "Tests were adjusted to actual selectors, and the accessibility issue was documented for future improvement."],
            ["Handling asynchronous loading and API errors.", "Users could click repeatedly or receive unclear feedback.", "Loading states, disabled buttons, and error popups were implemented."],
        ],
        font_size=7,
    )

    doc.add_heading("5.7 Code Explanation", 2)
    doc.add_heading("5.7.1 Route Configuration", 3)
    para(doc, "The route configuration in App.js is the entry point for frontend navigation. It maps each URL path to the correct React page component.")
    code(
        doc,
        "<Routes location={location}>\n"
        "  <Route path=\"/\" element={<Landing />} />\n"
        "  <Route path=\"/login\" element={<Login />} />\n"
        "  <Route path=\"/client/dashboard\" element={<ClientDashboard />} />\n"
        "  <Route path=\"/freelancer/jobs\" element={<FreelancerFindJobs />} />\n"
        "  <Route path=\"/freelancer/messages\" element={<Messages />} />\n"
        "</Routes>",
    )
    doc.add_heading("5.7.2 Login Validation", 3)
    para(doc, "Login validation is performed before calling the backend. If validation fails, errors are stored in state and rendered below the relevant fields.")
    code(
        doc,
        "const validateLogin = () => {\n"
        "  const errs = {};\n"
        "  if (!email || !email.includes('@')) errs.email = 'Please enter a valid email.';\n"
        "  if (!password) errs.password = 'Password is required.';\n"
        "  setErrors(errs);\n"
        "  return Object.keys(errs).length === 0;\n"
        "};",
    )
    doc.add_heading("5.7.3 Proposal Submission UI", 3)
    para(doc, "The freelancer Find Jobs page supports a real end-to-end workflow: browse jobs, open details, enter a cover letter and bid, then submit a proposal.")
    code(
        doc,
        "const payload = {\n"
        "  job_id: modalJob.id,\n"
        "  cover_letter: applyForm.cover_letter.trim(),\n"
        "  bid_amount: bid,\n"
        "  estimated_days: applyForm.estimated_days ? parseInt(applyForm.estimated_days, 10) : null,\n"
        "};\n"
        "await api.post('/proposals', payload);\n"
        "showToast('Application submitted!');",
    )
    add_img(doc, "testing/browser-screenshots/freelancer-e2e-proposal-confirmation.png", "Figure 5.3: Freelancer proposal workflow confirmation captured during browser E2E testing")


def chapter_6(doc):
    doc.add_heading("Chapter 6: System Testing", 1)
    doc.add_heading("6.1 Overview of Testing Process", 2)
    para(
        doc,
        "System testing for my frontend part focused on verifying that the React application worked as an actual user interface, not only as separate code files. The testing process included unit/component tests, frontend-backend integration checks, browser smoke tests, end-to-end workflow testing, build verification, route verification, and acceptance-readiness checks.",
    )
    para(
        doc,
        "Only results with actual evidence were marked as passed. Where a test required external human login or stakeholder approval, it was not falsely marked as passed.",
    )

    doc.add_heading("6.2 Testing Environment", 2)
    make_table(
        doc,
        ["Environment Area", "Details"],
        [
            ["Operating System", "Windows local development environment."],
            ["Frontend Server", "React development server at http://localhost:3000 / http://127.0.0.1:3000."],
            ["Backend Server", "FastAPI backend at http://127.0.0.1:8000."],
            ["API Base URL", "REACT_APP_API_URL=http://127.0.0.1:8000/api for local frontend testing."],
            ["Docker Services", "PostgreSQL, Redis, IPFS, and Hardhat local blockchain."],
            ["Health Evidence", "Backend /api/health returned 200 with database, Redis, IPFS, and blockchain ok; frontend root returned 200."],
            ["Unit Testing Tools", "Jest and React Testing Library with custom config for React Router v7 compatibility."],
            ["Browser Testing Tool", "Playwright Chromium used for browser smoke tests and freelancer end-to-end workflow."],
            ["Test Evidence Files", "frontend-jest-output.txt, frontend-integration-results.json, freeledger-browser-results.json, freeledger-freelancer-e2e-result.json, frontend-build-output.txt, browser screenshots, and browser-test-report.html."],
        ],
    )

    doc.add_heading("6.3 Test Cases", 2)
    make_table(
        doc,
        ["Test Case ID", "Testing Type", "Scenario", "Expected Result", "Actual Result", "Status"],
        [
            ["UT-001", "Unit Testing", "Landing page Sign in button navigation.", "Navigation should target /login.", "Jest confirmed Sign in navigation was called with /login.", "Pass"],
            ["UT-002", "Unit Testing", "Login form renders email and password inputs.", "Email and password inputs should appear.", "Jest confirmed both fields rendered.", "Pass"],
            ["UT-003", "Unit Testing", "Empty login form validation.", "Email and password validation messages should display.", "Both validation messages displayed.", "Pass"],
            ["UT-004", "Unit Testing", "Invalid email validation.", "Invalid email message should display.", "Please enter a valid email. displayed.", "Pass"],
            ["UT-005", "Unit Testing", "Valid login submission.", "Axios login should be called and dashboard navigation should occur.", "Mocked login stored tokens and navigated to client dashboard.", "Pass"],
            ["UT-006", "Unit Testing", "Google login button rendering.", "Google button should render when client ID is configured.", "Google script was simulated and button rendered.", "Pass"],
            ["INT-001", "Integration Testing", "Frontend login form to backend authentication and session storage.", "Valid login should redirect to dashboard.", "Browser test logged in and reached /client/dashboard.", "Pass"],
            ["INT-002", "Integration Testing", "Failed login from frontend to backend.", "Invalid credentials should show an error.", "Browser test displayed invalid credential error.", "Pass"],
            ["INT-003", "Integration Testing", "Authenticated API pages: jobs, contracts, profile, messages.", "Endpoints should return HTTP 200 with token.", "API checks returned HTTP 200.", "Pass"],
            ["SYS-001", "System Testing", "Frontend build.", "npm run build should compile.", "Build compiled successfully.", "Pass"],
            ["SYS-002", "System Testing", "Frontend and backend reachability.", "Frontend root and backend health should return HTTP 200.", "Both returned HTTP 200.", "Pass"],
            ["SYS-003", "System Testing", "Dashboard navigation.", "Authenticated user should reach Jobs, Contracts, Messages, and Profile.", "Playwright reached all listed routes from dashboard navigation.", "Pass"],
            ["SYS-004", "System Testing", "Responsive homepage.", "Homepage should render on desktop, tablet, and mobile viewports.", "Playwright confirmed branding rendered on all three viewports.", "Pass"],
            ["SYS-005", "System Testing", "Direct HTTP access to React routes.", "Routes should serve the React app shell.", "Command-line direct route checks returned HTTP 404 for several routes.", "Fail"],
            ["E2E-001", "End-to-End Testing", "Freelancer login â†’ browse jobs â†’ open details â†’ submit proposal â†’ receive confirmation.", "Freelancer should see proposal confirmation.", "Automated E2E completed and saw Application submitted confirmation.", "Pass"],
            ["ACC-001", "Acceptance Testing", "Registration, login, dashboards, jobs, profile, contracts, messaging.", "Core agreed frontend requirements should be covered.", "Automated/API/browser evidence covered the main frontend requirements.", "Pass"],
            ["ACC-002", "Acceptance Testing", "Real Google OAuth sign-in.", "User should sign in with real Google account.", "Blocked because real Google account consent was required during this cycle.", "Blocked"],
        ],
        font_size=7,
    )

    doc.add_heading("6.4 Test Results", 2)
    make_table(
        doc,
        ["Test Suite", "Total Test Cases", "Executed", "Passed", "Failed", "Blocked", "Not Executed", "Pass Percentage"],
        [
            ["Unit Testing", "10", "10", "10", "0", "0", "0", "100.00%"],
            ["Integration Testing", "11", "11", "11", "0", "0", "0", "100.00%"],
            ["System Testing", "14", "14", "11", "3", "0", "0", "78.57%"],
            ["End-to-End Testing", "1", "1", "1", "0", "0", "0", "100.00%"],
            ["Acceptance Testing", "4", "3", "2", "0", "1", "1", "66.67%"],
            ["Overall", "40", "39", "35", "3", "1", "1", "89.74%"],
        ],
        font_size=7,
    )

    doc.add_heading("6.5 Discussion of Failures", 2)
    para(
        doc,
        "The main frontend-related failure was direct HTTP access to React routes returning 404 in command-line checks. Routes such as /login, /client/dashboard, /client/messages, /freelancer/dashboard, and /freelancer/jobs were not consistently served as React app fallback routes when requested directly by HTTP tools. Browser client-side navigation worked during Playwright testing, which means the React router works after the app shell loads, but server fallback behavior needs improvement.",
    )
    para(
        doc,
        "A second issue found during testing was accessibility-related. The login form visually displays labels for email and password, but the labels are not programmatically associated with their inputs. This should be improved using htmlFor/id or aria-labelledby.",
    )
    para(
        doc,
        "The standard CRA npm test command also failed to discover tests because the workspace path included a hidden .gemini directory on Windows. This was treated as a testing-environment issue, not a product UI failure. A custom Jest configuration allowed the tests to execute successfully.",
    )

    doc.add_heading("6.6 Performance Testing", 2)
    para(
        doc,
        "Full load testing was outside the scope of my individual frontend report, but frontend performance-related observations were made. The production build compiled successfully, and browser smoke tests loaded the homepage, login page, dashboard navigation, and responsive views without visible timeout failures.",
    )
    make_table(
        doc,
        ["Performance / Build Check", "Observed Result"],
        [
            ["React production build", "Compiled successfully."],
            ["Main JavaScript bundle after gzip", "Approximately 249.83 kB."],
            ["Secondary JavaScript chunk after gzip", "Approximately 47.11 kB."],
            ["Main CSS after gzip", "Approximately 22.11 kB."],
            ["Browser homepage rendering", "Desktop, tablet, and mobile homepage smoke checks passed."],
        ],
    )

    doc.add_heading("6.7 Validation and Verification", 2)
    para(
        doc,
        "Validation was performed by comparing implemented frontend behavior with the functional expectations from the project: users should access the platform, register or log in, be redirected based on role, view dashboards, browse jobs, manage contracts/profile pages, use messaging routes, and submit proposals. Verification was performed through code inspection, build output, health checks, API integration checks, unit tests, and browser automation.",
    )
    para(
        doc,
        "The frontend met the main implemented requirements for local development. Improvements remain, especially route fallback behavior, full real Google OAuth verification, better accessibility attributes, and stronger deployment environment documentation.",
    )


def chapter_7(doc):
    doc.add_heading("Chapter 7: Conclusion and Critical Evaluation", 1)
    doc.add_heading("7.1 Summary of Work Done", 2)
    para(
        doc,
        "My frontend work in Capstone-II focused on turning the planned FreeLedger interface into a working React frontend. The completed frontend includes public landing page, authentication page, role-based client and freelancer dashboards, job browsing, proposal submission interface, contracts pages, profile pages, messages page, admin pages, central API service, Google authentication UI integration, responsive styles, and frontend testing evidence.",
    )
    para(
        doc,
        "The project changed significantly from the first report. Capstone-I described what the frontend should become. Capstone-II shows the actual implementation decisions, integration issues, and test results. The frontend is now not only a design concept but a working application layer connected to backend APIs and tested through unit and browser automation.",
    )

    doc.add_heading("7.2 Critical Appraisal of the Project", 2)
    doc.add_heading("7.2.1 Strengths", 3)
    for item in [
        "The frontend uses a clear React component structure, which makes the system easier to extend and maintain.",
        "Separate client and freelancer dashboards improve usability because each role receives a workspace related to their actual tasks.",
        "The central Axios service reduces duplicated API logic and keeps authentication headers consistent.",
        "Google authentication support improves accessibility for non-technical users who may not want to begin with wallet-based login.",
        "The UI has grown beyond static screens into real workflows such as login, dashboard navigation, job browsing, and proposal submission.",
        "Automated tests and browser screenshots provide stronger evidence than manual claims alone.",
    ]:
        bullet(doc, item)

    doc.add_heading("7.2.2 Weaknesses / Limitations", 3)
    for item in [
        "Direct HTTP access to some React routes still returned 404 during command-line checks, meaning route fallback needs to be configured more reliably.",
        "Some accessibility details, especially input-label association on the login form, need improvement.",
        "Real Google OAuth success was not fully executed in the test cycle because it required live Google sign-in and consent.",
        "The frontend depends heavily on correct .env configuration. If API URLs or CORS origins are missing, users see Network Error even when the UI itself is correct.",
        "Some originally planned blockchain/wallet interactions remained more complex to test than ordinary email/password flows.",
    ]:
        bullet(doc, item)

    doc.add_heading("7.3 Challenges Faced", 2)
    para(
        doc,
        "The most important challenge was frontend-backend integration. A frontend can look correct visually but still fail if the API base URL, backend server, CORS configuration, or authentication token handling is wrong. During testing, Network Error and CORS preflight issues showed how sensitive browser applications are to environment configuration.",
    )
    para(
        doc,
        "Another major challenge was testing. The standard npm test command failed to discover test files because of the Windows workspace path. React Router v7 also required additional Jest configuration and a TextEncoder polyfill. Instead of ignoring this, I created a custom test setup and documented the workaround clearly.",
    )
    para(
        doc,
        "A further challenge was separating frontend responsibility from backend/infrastructure problems. Python virtual environment errors and Windows blocking of DLL files were not frontend code bugs, but they still prevented frontend login testing because the backend had to be running.",
    )

    doc.add_heading("7.4 Future Improvements", 2)
    for item in [
        "Add proper React route fallback for all direct URLs in development and production hosting so direct navigation never returns 404.",
        "Improve form accessibility by connecting labels to inputs using htmlFor/id and adding aria attributes where necessary.",
        "Add stronger validation messages, such as separate messages for empty email and invalid email format.",
        "Add complete Google OAuth live testing using a valid test Google account and documented consent flow.",
        "Add more Playwright tests for freelancer contracts, client project posting, messaging thread creation, and admin workflows.",
        "Improve deployment documentation by including clear .env.example files for frontend and backend local setup.",
        "Add visual regression testing or screenshot comparison for key pages to prevent UI layout regressions.",
        "Improve mobile dashboard navigation and test all role pages on mobile, not only the homepage.",
    ]:
        bullet(doc, item)

    doc.add_heading("7.5 Personal Reflection", 2)
    para(
        doc,
        "Working on the frontend part of FreeLedger helped me understand that frontend development is not only about making pages look good. It also includes routing, validation, API communication, authentication state, role-based navigation, error handling, responsive design, accessibility, testing, and debugging environment issues. The biggest learning point for me was that user experience depends on both design and reliability.",
    )
    para(
        doc,
        "I also learned the importance of evidence-based reporting. Instead of writing that features worked without proof, I used build output, API checks, Jest tests, Playwright browser tests, screenshots, and documented defects. This made the Capstone-II report more authentic and helped separate passed, failed, blocked, and not-executed items properly.",
    )
    para(
        doc,
        "Overall, my frontend contribution changed FreeLedger from an early planned interface into a usable role-based web application. There are still improvements required, but the current frontend provides a stronger base for future development, testing, and deployment.",
    )

    doc.add_heading("7.6 Final Conclusion", 2)
    para(
        doc,
        "The Capstone-II frontend implementation successfully delivered the major frontend requirements for FreeLedger in local development. The application now includes a working public interface, authentication UI, role-specific dashboards, navigation, job/proposal workflow, contracts/profile/messages pages, Google authentication support, API integration, and automated testing evidence. The most important unresolved frontend issues are direct route fallback, accessibility improvements, and full live Google OAuth verification. Despite these limitations, the frontend has reached a functional stage and provides a strong foundation for completing and presenting the overall FreeLedger project.",
    )
    para(
        doc,
        "Note: This individual report is limited to the frontend development contribution of Bijee Dangol. Backend, database, blockchain, and infrastructure work are mentioned only where they affected frontend implementation, integration, or testing.",
    )


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    for style in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        doc.styles[style].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 3"].font.size = Pt(12)

    add_title(doc)
    chapter_5(doc)
    chapter_6(doc)
    chapter_7(doc)
    doc.save(WORKSPACE_OUT)
    doc.save(DOWNLOADS_OUT)
    print(WORKSPACE_OUT)
    print(DOWNLOADS_OUT)


if __name__ == "__main__":
    main()
